import argparse
from pathlib import Path


def set_run_fonts(run, font_name: str):
    run.font.name = font_name
    r = run._element.rPr
    rFonts = r.rFonts if r.rFonts is not None else r._add_rFonts()
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rFonts.set(qn("w:cs"), font_name)


def apply_compact_page_setup(document, margins_cm: float):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    m = Cm(margins_cm)
    section.top_margin = m
    section.bottom_margin = m
    section.left_margin = m
    section.right_margin = m


def apply_compact_styles(document, base_font: str, base_font_pt: int):
    normal = document.styles["Normal"]
    normal.font.name = base_font
    normal.font.size = Pt(base_font_pt)
    pf = normal.paragraph_format
    pf.line_spacing = 1.0
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    for style_name, size_pt in [
        ("Heading 1", base_font_pt + 3),
        ("Heading 2", base_font_pt + 2),
        ("Heading 3", base_font_pt + 1),
        ("Heading 4", base_font_pt),
    ]:
        try:
            st = document.styles[style_name]
        except KeyError:
            continue
        st.font.name = base_font
        st.font.size = Pt(size_pt)
        st.font.bold = True
        st.paragraph_format.space_before = Pt(2)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.0

    for style_name in ["List Bullet", "List Number"]:
        try:
            st = document.styles[style_name]
        except KeyError:
            continue
        st.font.name = base_font
        st.font.size = Pt(base_font_pt)
        st.paragraph_format.space_before = Pt(0)
        st.paragraph_format.space_after = Pt(0)
        st.paragraph_format.line_spacing = 1.0


def add_paragraph_compact(document, text: str, style: str = "Normal"):
    p = document.add_paragraph(text, style=style)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    return p


def md_to_docx(md_text: str, *, base_font: str, base_font_pt: int, margins_cm: float):
    document = Document()
    apply_compact_page_setup(document, margins_cm=margins_cm)
    apply_compact_styles(document, base_font=base_font, base_font_pt=base_font_pt)

    lines = [ln.rstrip("\n") for ln in md_text.splitlines()]
    pending_paragraph: list[str] = []

    def flush_paragraph():
        nonlocal pending_paragraph
        if not pending_paragraph:
            return
        text = " ".join([t.strip() for t in pending_paragraph if t.strip()])
        if text:
            add_paragraph_compact(document, text, style="Normal")
        pending_paragraph = []

    for raw in lines:
        line = raw.strip()
        if not line:
            flush_paragraph()
            continue

        if line.startswith("#"):
            flush_paragraph()
            level = len(line) - len(line.lstrip("#"))
            title = line[level:].strip()
            if not title:
                continue
            heading_style = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}.get(
                level, "Heading 4"
            )
            add_paragraph_compact(document, title, style=heading_style)
            continue

        if line.startswith(("-", "*")) and line[1:2] == " ":
            flush_paragraph()
            add_paragraph_compact(document, line[2:].strip(), style="List Bullet")
            continue

        if len(line) >= 3 and line[0].isdigit() and line[1:3] == ") ":
            flush_paragraph()
            add_paragraph_compact(document, line[3:].strip(), style="List Number")
            continue

        if line.startswith("- ") and line[2:].strip().startswith("`") and line[2:].strip().endswith("`"):
            flush_paragraph()
            add_paragraph_compact(document, line[2:].strip("`").strip(), style="List Bullet")
            continue

        pending_paragraph.append(line)

    flush_paragraph()

    for p in document.paragraphs:
        for run in p.runs:
            set_run_fonts(run, base_font)

    return document


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Convert Proposal_描述性改写版.md to a compact .docx (aiming for ~2 pages)."
    )
    parser.add_argument(
        "--in",
        dest="in_path",
        default=str(Path(__file__).resolve().parent.parent / "Proposal_描述性改写版.md"),
        help="Input Markdown path.",
    )
    parser.add_argument(
        "--out",
        dest="out_path",
        default=str(Path(__file__).resolve().parent / "Proposal_描述性改写版.docx"),
        help="Output .docx path.",
    )
    parser.add_argument(
        "--font",
        default="宋体",
        help="Base font name (default: 宋体).",
    )
    parser.add_argument(
        "--font-pt",
        type=int,
        default=10,
        help="Base font size in points (default: 10).",
    )
    parser.add_argument(
        "--margins-cm",
        type=float,
        default=1.5,
        help="Page margins in cm (default: 1.5). Smaller = more compact.",
    )
    args = parser.parse_args()

    in_path = Path(args.in_path)
    out_path = Path(args.out_path)

    if not in_path.exists():
        raise SystemExit(f"Input file not found: {in_path}")

    md_text = in_path.read_text(encoding="utf-8")
    doc = md_to_docx(
        md_text,
        base_font=args.font,
        base_font_pt=args.font_pt,
        margins_cm=args.margins_cm,
    )

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    print(
        "Tip: Page count depends on your Word defaults/printer. If it's >2 pages, try --font-pt 9 and/or --margins-cm 1.2."
    )
    return 0


if __name__ == "__main__":
    try:
        from docx import Document  # type: ignore
        from docx.shared import Cm, Mm, Pt  # type: ignore
        from docx.oxml.ns import qn  # type: ignore
    except Exception as e:  # pragma: no cover
        raise SystemExit(
            "Missing dependency python-docx. Install with: python -m pip install python-docx\n"
            f"Error: {e}"
        )

    raise SystemExit(main())

